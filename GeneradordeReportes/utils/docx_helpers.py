from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .helpers import get_region_language
from .text_templates import text_templates


def render_title(doc: Document, country_code: str, year: int):
    """Inserta título centrado en el documento"""
    lang = get_region_language(country_code)
    text = text_templates["title"][lang].format(year=year)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def render_intro_and_table(
    doc: Document,
    country_code: str,
    farmer_name: str,
    datos: dict,  # Ej: {"farmercode":"CR-001", "contractcode":"CR0092", "planting_year":2025, "contract_trees":1000}
    code: str = None
):
    """Genera la sección de introducción con tabla 4×2"""
    lang = get_region_language(country_code)

    # --- INTRO: convierte <br> en párrafos ---
    intro_tpl = text_templates["intro"][lang]
    intro_text = intro_tpl.format(farmername=farmer_name)
    # Creamos la tabla 4×2
    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    # Merge de columna izquierda
    left = table.cell(0, 0).merge(table.cell(3, 0))
    # Añadimos solo los párrafos útiles, cortando por doble salto
    for section in [s.strip() for s in intro_text.split("<br><br>") if s.strip()]:
        left.add_paragraph(section)

    # --- CELDAS DERECHA con placeholders de text_templates ---
    for i, key in enumerate(("farmercode", "contractcode", "planting_year", "contract_trees")):
        tpl = text_templates["cells_right"][key][lang]
        text = tpl.format(**datos)
        cell = table.cell(i, 1)
        cell.text = text

    doc.add_paragraph()
    return table
