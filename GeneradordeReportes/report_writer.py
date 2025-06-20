#GeneradodeReportes/report_writer

from core.libs import pd, os, text
from core.db import get_engine

from docx import Document
from docx.shared import Pt, Inches
from GeneradordeReportes.utils.dynamic_text_blocks import fetch_dynamic_values, format_paragraphs
from GeneradordeReportes.utils.docx_helpers import render_title, render_intro_and_table
from GeneradordeReportes.graficadorG1Mortalidad import generar_mortalidad
from GeneradordeReportes.graficadorG2Altura import generar_altura
from GeneradordeReportes.graficadorG3Crecimiento import generar_crecimiento
from GeneradordeReportes.graficadorG4DefectosyPlagas import generar_tabla_sanidad
import argparse
from GeneradordeReportes.utils.helpers import get_region_language
from GeneradordeReportes.utils.text_templates import text_templates
from GeneradordeReportes.utils.text_calculations import get_mortality_metrics
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from GeneradordeReportes.utils.config import EXPORT_WIDTH_INCHES, EXPORT_HEIGHT_INCHES

# Configuración de rutas
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(BASE_DIR, "GeneradordeReportes", "templates", "Template.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "GeneradordeReportes", "outputs", "reports")
IMG_TMP_DIR = os.path.join(BASE_DIR, "GeneradordeReportes", "temp_images")

# Crear directorios necesarios
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(IMG_TMP_DIR, exist_ok=True)


def crear_reporte(code: str, country: str, year: int, engine) -> str:
    """Función principal para generar el reporte completo"""

    # 1. Generación de gráficas
    output_root = os.path.join(BASE_DIR, "GeneradordeReportes", "outputs")
    paths = {
        'G1': generar_mortalidad(code, country, year, engine, output_root=output_root),
        'G2': generar_altura(code, country, year, engine, output_root=output_root),
        'G3': generar_crecimiento(code, country, year, engine, output_root=output_root),
    }
    metrics = get_mortality_metrics(engine, country, year, code)

    # === 🚩 Agrega este bloque ===
    if metrics["dead"] + metrics["alive"] == 0:
        print(f"⏩ Contrato {code} sin árboles censados. Se omite.")
        return  # Nada que reportar

    if metrics["rate"] == 100:
        print(f"⏩ Contrato {code} con 100% de mortalidad. Se omite.")
        return  # Todo muerto, omitimos reporte
    # ==============================

    # 2. Validación de gráficas generadas
    for key, path in paths.items():
        if not path or not os.path.isfile(path):
            print(f"⚠️ Gráfica {key} no generada para {code}: {path}")
            # path = None  # (opcional) explícitamente marcarlo como None

    # 3. Generar tabla de sanidad
    df_sanidad = generar_tabla_sanidad(code, country, year, engine)

    # 4. Crear documento Word
    doc = Document(TEMPLATE)

    # 5. Sección de título
    render_title(doc, country, year)

    # 6. Sección introductoria: valores dinámicos y datos de contrato
    values = fetch_dynamic_values(code, country, year)
    datos = {
        "contractcode": values.get("contractcode", code),
        "farmer_number": values.get("farmer_number", ''),
        "planting_year": values.get("planting_year", ''),
        "contract_trees": values.get("contract_trees", ''),
    }

    # —— Obtener nombre real del productor ——
    with engine.connect() as conn:
        sql = text(
            'SELECT contract_name FROM masterdatabase.contract_farmer_information WHERE contract_code = :fcode'
        )
        producer_name = conn.execute(sql, {'fcode': datos['contractcode']}).scalar_one()

    # Render de la sección introductoria
    render_intro_and_table(
        doc,
        country,
        farmer_name=producer_name,
        datos=datos,
        code=code
    )

    # — 8) Insertar solo las imágenes G1–G3 con encabezados y contenidos dinámicos —
    resumen_dir = os.path.join(BASE_DIR, "GeneradordeReportes", "outputs", code, "Resumen")
    grafs = [
        ("G1", "G1_Mortality_"),
        ("G2", "G2_Altura_"),
        ("G3", "G3_Crecimiento_"),
    ]
    lang = get_region_language(country)

    for key, prefix in grafs:
        img_path = os.path.join(resumen_dir, f"{prefix}{code}.png")
        if not os.path.exists(img_path):
            print(f"⚠️ Gráfica {key} no encontrada: {img_path}")
            continue

        if key == "G1":
            # Título dinámico de mortalidad
            mort_title = (
                f"{text_templates['chart_titles']['mortality'][lang]}: "
                f"{metrics['rate']:.1f}%"
            )
            doc.add_heading(mort_title, level=2)

            # Crear tabla 1x2 sin bordes
            table = doc.add_table(rows=1, cols=2)
            tbl = table._tbl
            tblPr = tbl.tblPr
            for node in tblPr.findall(qn('w:tblBorders')):
                tblPr.remove(node)
            borders = OxmlElement('w:tblBorders')
            for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                elm = OxmlElement(f'w:{tag}')
                elm.set(qn('w:val'), 'nil')
                borders.append(elm)
            tblPr.append(borders)

            # Contenido celda izquierda
            cell_text, cell_img = table.rows[0].cells
            section2_title = text_templates["section_headers"]["G2"][lang]
            mort_text = text_templates["mortality_text"][lang].format(
                dead_per_100=metrics['dead_per_100'],
                alive=metrics['survivors_estimated']
            )

            # Agrega texto completo en la celda izquierda
            cell_text.paragraphs[0].add_run(mort_text)
            cell_text.add_paragraph(section2_title, style="Heading 2")
            for paragraph in format_paragraphs(values, country):
                cell_text.add_paragraph(paragraph)

            # Imagen G1 en la celda derecha
            run = cell_img.paragraphs[0].add_run()
            run.add_picture(img_path, width=Inches(8.5 / 2.54), height=Inches(5.8 / 2.54))

            # Después de la tabla, insertar gráficas G2 y G3
            g2_path = os.path.join(resumen_dir, f"G2_Altura_{code}.png")
            g3_path = os.path.join(resumen_dir, f"G3_Crecimiento_{code}.png")

            if os.path.exists(g2_path) and os.path.exists(g3_path):
                table_imgs = doc.add_table(rows=2, cols=1)
                table_imgs.autofit = True  # Opcional
                # Elimina bordes (opcional)
                tbl = table_imgs._tbl
                for border in tbl.xpath(".//w:tblBorders"):
                    tbl.remove(border)
                # G2
                run_g2 = table_imgs.cell(0, 0).paragraphs[0].add_run()
                run_g2.add_picture(g2_path, width=Inches(EXPORT_WIDTH_INCHES))
                # G3
                run_g3 = table_imgs.cell(1, 0).paragraphs[0].add_run()
                run_g3.add_picture(g3_path, width=Inches(EXPORT_WIDTH_INCHES))
            elif os.path.exists(g2_path):
                doc.add_picture(g2_path, width=Inches(EXPORT_WIDTH_INCHES))
            elif os.path.exists(g3_path):
                doc.add_picture(g3_path, width=Inches(EXPORT_WIDTH_INCHES))

            continue  # Saltar a siguiente


        elif key == "G2":

            continue


        elif key == "G3":

            continue

        # Inserta la imagen y luego salta de página
        doc.add_picture(img_path)

    # 8. Tabla de sanidad
    if df_sanidad is not None and not df_sanidad.empty:
        doc.add_heading('Distribución de Plagas, Defectos y Enfermedades', level=2)
        table = doc.add_table(rows=1, cols=len(df_sanidad.columns), style='Table Grid')
        hdr_cells = table.rows[0].cells
        for idx, col_name in enumerate(df_sanidad.columns):
            run = hdr_cells[idx].paragraphs[0].add_run(str(col_name))
            run.bold = True
            run.font.size = Pt(10)
        for _, row in df_sanidad.iterrows():
            cells = table.add_row().cells
            for idx, val in enumerate(row):
                cells[idx].text = str(val)
    else:
        print(f"⚠️ No hay datos de sanidad para {code}")

    #9. Notas y observaciones
    lang = get_region_language(country)
    doc.add_heading(text_templates["section_headers"]["Notas"][lang], level=2)
    doc.add_paragraph("")  # Un enter (párrafo vacío)
    doc.add_paragraph("")  # Otro enter

    doc.add_heading(text_templates["section_headers"]["Recomendaciones"][lang], level=2)

    # 10. Guardar documento
    import time
    def guardar_documento(doc, out_path: str, intentos=3):
        for i in range(intentos):
            try:
                doc.save(out_path)
                return True
            except PermissionError:
                print(f"⚠️ Error de permisos (intento {i+1}/{intentos}). Cerrando Word...")
                time.sleep(2)
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Quit()
                except:
                    pass
        return False

    out_name = f"Reporte_{code}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    if guardar_documento(doc, out_path):
        print(f"✅ Reporte creado: {out_path}")
    else:
        print(f"❌ Error crítico: No se pudo guardar {out_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('code')
    parser.add_argument('--country', '-c', required=True)
    parser.add_argument('--year', '-y', type=int, required=True)
    args = parser.parse_args()

    engine = get_engine()
    crear_reporte(args.code, args.country, args.year, engine)


