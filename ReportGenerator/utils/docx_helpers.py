from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .helpers import get_region_language
from .text_templates import text_templates


def render_title(doc: Document, country_code: str, year: int):
    """Inserta título centrado en el documento"""
    lang = get_region_language(country_code)
    text = text_templates["title"][lang].format(year=year)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _remove_table_borders(table):
    """Elimina todos los bordes de la tabla."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # quita cualquier tblBorders existente
    for node in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(node)
    # añade tblBorders con valor 'nil' en todas las aristas
    borders = OxmlElement('w:tblBorders')
    for tag in ('top','left','bottom','right','insideH','insideV'):
        elm = OxmlElement(f'w:{tag}')
        elm.set(qn('w:val'), 'nil')
        borders.append(elm)
    tblPr.append(borders)

def render_intro_and_table(
    doc: Document,
    country_code: str,
    farmer_name: str,
    datos: dict,  # Ej: {"farmercode":"CR-001", "contractcode":"CR0092", "planting_year":2025, "contract_trees":1000}
    code: str = None
):
    """Genera la sección de introducción con tabla 4×2"""
    lang = get_region_language(country_code)
    intro_tpl = text_templates["intro"][lang]
    # formateo de <br><br>
    sections = [s.strip() for s in intro_tpl.format(farmername=farmer_name).split("<br><br>") if s.strip()]

    ## tabla 4×2 sin bordes
    table = doc.add_table(rows=4, cols=2)
    _remove_table_borders(table)

    # merge a una sola celda izquierda
    left = table.cell(0, 0).merge(table.cell(3, 0))
    # en lugar de add_paragraph sin más, metemos el primer párrafo en el slot 0
    for i, text in enumerate(sections):
        if i == 0:
            left.paragraphs[0].text = text
        else:
            left.add_paragraph(text)

    # derecha: los 4 datos
    for i, key in enumerate(("farmercode", "contractcode", "planting_year", "contract_trees")):
        tpl = text_templates["cells_right"][key][lang]
        txt = tpl.format(**datos)
        cell = table.cell(i, 1)
        cell.text = txt

    doc.add_paragraph()
    return table
